from django.shortcuts import render, redirect, get_object_or_404
import threading
from django.core.mail import send_mail, EmailMultiAlternatives
from django.conf import settings
from django.contrib.auth import authenticate, login, logout
from django.contrib.auth.models import User
from django.contrib import messages
from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
from django.views.decorators.http import require_http_methods
from django.utils import timezone
from datetime import datetime, timedelta
from django.db.models import Count
from django.db.models.functions import TruncDay, TruncMonth
import json
from .models import Admin, Student, Popup, StudentDocument, ApplicationDocument, Message, AdminLog
from django.db import models
from home.models import Program, Application
from django.http import HttpResponse
import csv
import os
from django.utils.html import strip_tags
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def landing_page_view(request):
    try:
        # Fetch active programs
        active_programs = Program.objects.filter(is_active=True)
    except Exception:
        # Fallback if migration hasn't run yet or other error
        active_programs = []
    
    return render(request, 'accounts/landing_page.html', {'active_programs': active_programs})


def login_view(request):
    if request.method == 'POST':
        username = request.POST['username']
        password = request.POST['password']

        # Check if user is Admin
        admin_user = Admin.objects.filter(admin_name=username, password=password).first()
        if admin_user:
            request.session['user_role'] = 'admin'
            request.session['user_id'] = admin_user.admin_id
            
            # Log the login action
            AdminLog.objects.create(admin=admin_user, action="Signed in")

            return redirect('accounts:admin_dashboard')
        # Check if user is Student
        student_user = Student.objects.filter(username=username, password=password).first()
        if student_user:
            # Check if student is approved
            if student_user.status == 'active':
                request.session['user_role'] = 'student'
                request.session['user_id'] = student_user.pk
                request.session['username'] = student_user.username
                return redirect('accounts:student_dashboard')
            elif student_user.status == 'pending':
                return render(request, 'accounts/login.html', {'error': 'Your account is pending approval. Please wait for admin approval.'})
            elif student_user.status == 'rejected':
                return render(request, 'accounts/login.html', {'error': 'Your account has been rejected. Please contact the administrator.'})
            elif student_user.status == 'inactive':
                return render(request, 'accounts/login.html', {'error': 'Account disabled'})

        # Invalid credentials
        return render(request, 'accounts/login.html', {'error': 'Invalid username or password'})

    return render(request, 'accounts/login.html')

import re

def register_view(request):
    if request.method == 'POST':
        username = request.POST.get('username')
        password = request.POST.get('password')
        confirm_password = request.POST.get('confirm_password')
        
        # Password Validation
        if not re.match(r'^(?=.*\d)(?=.*[A-Z]).{8,}$', password):
            messages.error(request, "Password must be at least 8 characters long, contain at least one capital letter and one number.")
            return render(request, 'accounts/register.html')

        if password != confirm_password:
             messages.error(request, "Passwords do not match.")
             return render(request, 'accounts/register.html')
        
        first_name = request.POST.get('first_name')
        last_name = request.POST.get('last_name')
        email = request.POST.get('email')

        # Check for duplicates
        if Student.objects.filter(email=email).exists():
            messages.error(request, "Email address is already in use.")
            return render(request, 'accounts/register.html')
            
        if Student.objects.filter(username=username).exists():
            messages.error(request, "Username is already taken.")
            return render(request, 'accounts/register.html')

        bday = request.POST.get('bday')
        address = request.POST.get('address')
        barangay = request.POST.get('barangay')
        student_type = request.POST.get('student_type')
        contact_num = request.POST.get('contact_num')
        program_and_yr = request.POST.get('program_and_yr')
        sex = request.POST.get('sex')
        current_school = request.POST.get('current_school')
        documents = request.FILES.getlist('document')
        
        if student_type != 'Undergraduate':
            program_and_yr = None

        first_doc = documents[0] if documents else None

        student = Student.objects.create(
            username=username,
            password=password,
            first_name=first_name,
            last_name=last_name,
            email=email,
            bday=bday,  
            address=address,
            barangay=barangay,
            student_type=student_type,
            contact_num=contact_num,
            program_and_yr=program_and_yr,
            sex=sex,
            current_school=current_school,
            doc_submitted=first_doc,
            status='pending' 
        )
        
        for doc in documents:
            StudentDocument.objects.create(student=student, file=doc)

        student.save()
        
        # Clear google data from session
        messages.success(request, "Registration successful! Your account is pending admin approval. You will be notified once approved.")
        return redirect('accounts:login')

    return render(request, 'accounts/register.html')

def _send_email_async(subject, message, recipient_list):
    """Helper to send email in a separate thread."""
    def _send():
        send_mail(subject, message, settings.DEFAULT_FROM_EMAIL, recipient_list, fail_silently=False)
    threading.Thread(target=_send).start()

def approve_student(request, student_id):
    if request.method == 'POST':
        student = get_object_or_404(Student, pk=student_id)
        action = request.POST.get('action')
        
        if action == 'approve':
            student.status = 'active'
            student.save()
            # Send approval email
            subject = 'Your application has been approved'
            message = f'Dear {student.first_name},\n\nYour application has been approved. You can now log in to your account.'
            _send_email_async(subject, message, [student.email])
        elif action == 'reject':
            student.status = 'rejected'
            student.save()
            # Send rejection email
            subject = 'Your application has been rejected'
            message = f'Dear {student.first_name},\n\nWe regret to inform you that your application was rejected. Please contact the administration for more details.'
            _send_email_async(subject, message, [student.email])
        else:
            return JsonResponse({'status': 'error', 'message': 'Invalid action'})
        return JsonResponse({'status': 'success'})


def admin_dashboard(request):
    admin_id = request.session.get('user_id')
    if not admin_id:
        return redirect('login')
    admin = Admin.objects.get(admin_id=admin_id)    
    programs = Program.objects.all()
    
    admins_list = Admin.objects.all()
    admin_logs = AdminLog.objects.select_related('admin').order_by('-timestamp')[:50]

    return render(request, 'accounts/admin_dashboard.html', {
        'admin': admin, 
        'programs': programs,
        'admins_list': admins_list,
        'admin_logs': admin_logs
    })


def student_dashboard(request):
    student_id = request.session.get('user_id')
    if not student_id:
        return redirect('login')
    student = get_object_or_404(Student, pk=student_id)
    return render(request, 'accounts/student_dashboard.html', {'student': student})

def logout_view(request):
    request.session.flush()  # clears all session data
    return redirect('accounts:login')


def create_program(request):
    print("🟢 create_program reached:", request.method)
    if request.method == 'POST':
        admin_id = request.session.get('user_id')
        program_name = request.POST.get('program_name')
        requirements = request.POST.get('requirements')
        document_requirements = request.POST.getlist('document_requirements')
        print("📦 POST DATA:", request.POST)

        program_type = request.POST.get('program_type')

        if not program_name:
            messages.error(request, "Program name is required.")
            return redirect('accounts:admin_dashboard')

        Program.objects.create(program_name=program_name, requirements=requirements, document_requirements=document_requirements, program_type=program_type)
        
        if admin_id:
            current_admin = Admin.objects.get(admin_id=admin_id)
            AdminLog.objects.create(admin=current_admin, action=f"Created program '{program_name}'")

        messages.success(request, f"Program '{program_name}' created successfully!")
        print("✅ Program created:", program_name)
        return redirect('accounts:admin_dashboard')

    messages.error(request, "Invalid request.")
    print("❌ Invalid request method")
    return redirect('accounts:admin_dashboard')


@csrf_exempt
@require_http_methods(["POST"])
def create_student_application(request):
    """Student submits application to a program"""
    student_id = request.session.get('user_id')
    if not student_id:
        return JsonResponse({'success': False, 'error': 'Not authenticated'}, status=401)

    try:
        program_id = request.POST.get('program') or request.POST.get('program_id')
        motivation = request.POST.get('motivation', '')

        if not program_id:
            return JsonResponse({'success': False, 'error': 'Program is required'}, status=400)

        # Get entities
        student = get_object_or_404(Student, pk=student_id)
        program = get_object_or_404(Program, program_id=program_id)

        app = Application.objects.create(
            student=student,
            program=program,
            requirement_status='submitted',
            remarks=motivation,
        )

        # Handle multiple supporting documents (new uploads)
        documents = request.FILES.getlist('supporting_docs')
        for doc in documents:
            ApplicationDocument.objects.create(application=app, file=doc)

        # Handle imported documents from Student's uploaded files
        imported_doc_ids = request.POST.getlist('imported_docs')
        if imported_doc_ids:
            from .models import StudentDocument # ensure imported
            for doc_id in imported_doc_ids:
                try:
                    student_doc = StudentDocument.objects.get(id=doc_id, student=student)
                    if student_doc.file:
                        # Create an ApplicationDocument referencing the same file
                        ApplicationDocument.objects.create(
                            application=app, 
                            file=student_doc.file
                        )
                except StudentDocument.DoesNotExist:
                    pass # Skip invalid or unauthorized document IDs

        return JsonResponse({'success': True, 'application_id': app.app_id})
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=500)
# Popup Management Views
def admin_stats(request):
    """Get admin dashboard statistics"""
    admin_id = request.session.get('user_id')
    if not admin_id:
        return JsonResponse({'error': 'Not authenticated'}, status=401)
    
    try:
        total_students = Student.objects.filter(status='active').count()
        active_popups = Popup.objects.filter(is_active=True).count()
        pending_student_registrations = Student.objects.filter(status='pending').count()
        pending_program_applications = Application.objects.filter(requirement_status='submitted').count()
        total_programs = Program.objects.count()
        
        # New: Unread student messages for inbox badge
        unread_messages = Message.objects.filter(is_read=False, sender_type='student').count()
        
        return JsonResponse({
            'total_students': total_students,
            'active_popups': active_popups,
            'pending_applications': pending_student_registrations, # Keeping for backward compatibility
            'pending_student_registrations': pending_student_registrations,
            'pending_program_applications': pending_program_applications,
            'total_programs': total_programs,
            'unread_messages': unread_messages
        })

    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)


@csrf_exempt
@require_http_methods(["POST"])
def admin_change_password(request):
    """Change admin password"""
    admin_id = request.session.get('user_id')
    if not admin_id:
        return JsonResponse({'error': 'Not authenticated'}, status=401)
    
    try:
        data = json.loads(request.body)
        new_password = data.get('new_password')
        confirm_password = data.get('confirm_password')
        
        if not new_password or not confirm_password:
            return JsonResponse({'success': False, 'error': 'Both password fields are required.'})
            
        if new_password != confirm_password:
            return JsonResponse({'success': False, 'error': 'Passwords do not match.'})
            
        # Password strength validation (optional but good to have)
        # Removed per user request
        
        admin = Admin.objects.get(admin_id=admin_id)
        admin.password = new_password
        admin.save()
        
        # Log the password change action
        AdminLog.objects.create(admin=admin, action="Changed password")
        
        return JsonResponse({'success': True, 'message': 'Password changed successfully.'})
        
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=500)


@csrf_exempt
@require_http_methods(["POST"])
def create_admin(request):
    """Create a new admin"""
    admin_id = request.session.get('user_id')
    if not admin_id:
        return JsonResponse({'error': 'Not authenticated'}, status=401)
    
    try:
        data = json.loads(request.body)
        new_username = data.get('username')
        new_full_name = data.get('full_name', '')
        new_password = data.get('password')
        confirm_password = data.get('confirm_password')

        if not new_username or not new_password or not confirm_password:
            return JsonResponse({'success': False, 'error': 'Username and password fields are required.'})

        if new_password != confirm_password:
            return JsonResponse({'success': False, 'error': 'Passwords do not match.'})

        if Admin.objects.filter(admin_name=new_username).exists():
            return JsonResponse({'success': False, 'error': 'An admin with this username already exists.'})

        new_admin = Admin.objects.create(
            admin_name=new_username,
            full_name=new_full_name,
            password=new_password
        )

        current_admin = Admin.objects.get(admin_id=admin_id)
        AdminLog.objects.create(admin=current_admin, action=f"Created new admin: {new_admin.admin_name}")

        return JsonResponse({'success': True, 'message': 'Admin created successfully.'})

    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=500)


def get_popups(request):
    """Get all popups for the admin"""
    admin_id = request.session.get('user_id')
    if not admin_id:
        return JsonResponse({'error': 'Not authenticated'}, status=401)
    
    try:
        popups = Popup.objects.all().order_by('-created_at')
        popup_data = []
        
        for popup in popups:
            popup_data.append({
                'id': popup.id,
                'title': popup.title,
                'message': popup.message,
                'popup_type': popup.popup_type,
                'is_active': popup.is_active,
                'created_at': popup.created_at.isoformat(),
                'updated_at': popup.updated_at.isoformat(),
                'expires_at': popup.expires_at.isoformat() if popup.expires_at else None
            })
        
        return JsonResponse({'popups': popup_data})
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)


def get_popup(request, popup_id):
    """Get a specific popup"""
    admin_id = request.session.get('user_id')
    if not admin_id:
        return JsonResponse({'error': 'Not authenticated'}, status=401)
    
    try:
        popup = get_object_or_404(Popup, id=popup_id)
        return JsonResponse({
            'id': popup.id,
            'title': popup.title,
            'message': popup.message,
            'popup_type': popup.popup_type,
            'is_active': popup.is_active,
            'created_at': popup.created_at.isoformat(),
            'updated_at': popup.updated_at.isoformat(),
            'expires_at': popup.expires_at.isoformat() if popup.expires_at else None
        })
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)


def get_student_popups(request):
    """Get active popups for students"""
    student_id = request.session.get('user_id')
    if not student_id:
        return JsonResponse({'error': 'Not authenticated'}, status=401)
    
    try:
        now = timezone.now()
        student = get_object_or_404(Student, pk=student_id)

        # Fetch active popups that haven't expired and not seen by student
        popups = Popup.objects.filter(
            is_active=True
        ).filter(
            models.Q(expires_at__isnull=True) | models.Q(expires_at__gt=now)
        ).exclude(
            seen_by=student
        ).order_by('-created_at')
        
        popup_data = []
        for popup in popups:
            popup_data.append({
                'id': popup.id,
                'title': popup.title,
                'message': popup.message,
                'popup_type': popup.popup_type,
                'created_at': popup.created_at.isoformat()
            })
        
        return JsonResponse({'popups': popup_data})
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)


@csrf_exempt
@require_http_methods(["POST"])
def mark_popup_viewed(request, popup_id):
    """Mark a popup as viewed by the student"""
    student_id = request.session.get('user_id')
    if not student_id:
        return JsonResponse({'error': 'Not authenticated'}, status=401)
    
    try:
        popup = get_object_or_404(Popup, id=popup_id)
        student = get_object_or_404(Student, pk=student_id)
        
        popup.seen_by.add(student)
        
        return JsonResponse({'success': True})
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=500)


@csrf_exempt
@require_http_methods(["POST"])
def create_popup(request):
    """Create a new popup"""
    admin_id = request.session.get('user_id')
    if not admin_id:
        return JsonResponse({'error': 'Not authenticated'}, status=401)
    
    try:
        data = json.loads(request.body)
        
        # Parse expires_at if provided
        expires_at = None
        if data.get('expires_at'):
            expires_at = datetime.fromisoformat(data['expires_at'].replace('Z', '+00:00'))
        
        popup = Popup.objects.create(
            title=data['title'],
            message=data['message'],
            popup_type=data['popup_type'],
            is_active=data.get('is_active', True),
            expires_at=expires_at
        )
        
        current_admin = Admin.objects.get(admin_id=admin_id)
        AdminLog.objects.create(admin=current_admin, action=f"Created popup '{popup.title}'")
        
        return JsonResponse({
            'success': True,
            'popup_id': popup.id,
            'message': 'Popup created successfully'
        })
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=500)


@csrf_exempt
@require_http_methods(["PUT"])
def edit_popup(request, popup_id):
    """Edit an existing popup"""
    admin_id = request.session.get('user_id')
    if not admin_id:
        return JsonResponse({'error': 'Not authenticated'}, status=401)
    
    try:
        popup = get_object_or_404(Popup, id=popup_id)
        data = json.loads(request.body)
        
        # Parse expires_at if provided
        expires_at = None
        if data.get('expires_at'):
            expires_at = datetime.fromisoformat(data['expires_at'].replace('Z', '+00:00'))
        
        popup.title = data['title']
        popup.message = data['message']
        popup.popup_type = data['popup_type']
        popup.is_active = data.get('is_active', True)
        popup.expires_at = expires_at
        popup.save()
        
        current_admin = Admin.objects.get(admin_id=admin_id)
        AdminLog.objects.create(admin=current_admin, action=f"Edited popup '{popup.title}'")
        
        return JsonResponse({
            'success': True,
            'message': 'Popup updated successfully'
        })
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=500)


@csrf_exempt
@require_http_methods(["POST"])
def toggle_popup(request, popup_id):
    """Toggle popup active status"""
    admin_id = request.session.get('user_id')
    if not admin_id:
        return JsonResponse({'error': 'Not authenticated'}, status=401)
    
    try:
        popup = get_object_or_404(Popup, id=popup_id)
        popup.is_active = not popup.is_active
        popup.save()
        
        action_word = "Activated" if popup.is_active else "Deactivated"
        current_admin = Admin.objects.get(admin_id=admin_id)
        AdminLog.objects.create(admin=current_admin, action=f"{action_word} popup '{popup.title}'")
        
        return JsonResponse({
            'success': True,
            'is_active': popup.is_active,
            'message': f'Popup {"activated" if popup.is_active else "deactivated"} successfully'
        })
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=500)


@csrf_exempt
@require_http_methods(["POST"])
def delete_popup(request, popup_id):
    """Delete a popup"""
    admin_id = request.session.get('user_id')
    if not admin_id:
        return JsonResponse({'error': 'Not authenticated'}, status=401)
    
    try:
        popup = get_object_or_404(Popup, id=popup_id)
        popup_title = popup.title
        popup.delete()
        
        current_admin = Admin.objects.get(admin_id=admin_id)
        AdminLog.objects.create(admin=current_admin, action=f"Deleted popup '{popup_title}'")
        
        return JsonResponse({
            'success': True,
            'message': 'Popup deleted successfully'
        })
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=500)


# Student Application Management Views
def get_student_applications(request):
    """Get all student applications for admin review"""
    admin_id = request.session.get('user_id')
    if not admin_id:
        return JsonResponse({'error': 'Not authenticated'}, status=401)
    
    try:
        students = Student.objects.all().order_by('-student_id')
        application_data = []
        
        for student in students:
            student_docs = student.documents.all()
            doc_list = []
            for d in student_docs:
                doc_list.append({
                    'id': d.id,
                    'name': d.document_name,
                    'url': d.file.url if d.file else None
                })
                
            application_data.append({
                'id': student.pk,
                'username': student.username,
                'first_name': student.first_name,
                'last_name': student.last_name,
                'email': student.email,
                'bday': student.bday.isoformat(),
                'address': student.address,
                'barangay': student.barangay,
                'student_type': student.student_type,
                'contact_num': student.contact_num,
                'program_and_yr': student.program_and_yr,
                'status': student.status,
                'doc_submitted': student.doc_submitted.url if student.doc_submitted else None,
                'student_documents': doc_list,
                'created_at': student.created_at.isoformat() if student.created_at else None
            })
        
        return JsonResponse({'applications': application_data})
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)


@csrf_exempt
@require_http_methods(["POST"])
def approve_student(request, student_id):
    """Approve a student application"""
    admin_id = request.session.get('user_id')
    if not admin_id:
        return JsonResponse({'error': 'Not authenticated'}, status=401)
    
    try:
        student = get_object_or_404(Student, pk=student_id)
        student.status = 'active'
        student.save()
        
        current_admin = Admin.objects.get(admin_id=admin_id)
        AdminLog.objects.create(admin=current_admin, action=f"Approved student {student.username}")
        
        return JsonResponse({
            'success': True,
            'message': 'Student application approved successfully'
        })
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=500)


@csrf_exempt
@require_http_methods(["POST"])
def reject_student(request, student_id):
    """Reject a student application"""
    admin_id = request.session.get('user_id')
    if not admin_id:
        return JsonResponse({'error': 'Not authenticated'}, status=401)
    
    try:
        student = get_object_or_404(Student, pk=student_id)
        student.status = 'rejected'
        student.save()
        
        current_admin = Admin.objects.get(admin_id=admin_id)
        AdminLog.objects.create(admin=current_admin, action=f"Rejected student {student.username}")
        
        return JsonResponse({
            'success': True,
            'message': 'Student application rejected successfully'
        })
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=500)


@csrf_exempt
@require_http_methods(["POST"])
def toggle_student_status(request, student_id):
    """Toggle student active status"""
    admin_id = request.session.get('user_id')
    if not admin_id:
        return JsonResponse({'error': 'Not authenticated'}, status=401)
    
    try:
        student = get_object_or_404(Student, pk=student_id)
        
        # If currently active, deactivate (set to inactive/suspended)
        # Using 'inactive' as the status for deactivated students
        if student.status == 'active':
            student.status = 'inactive'
            action = 'deactivated'
            log_action = "Deactivated"
        else:
            # If rejected, pending, or inactive, set to active
            student.status = 'active'
            action = 'activated'
            log_action = "Activated"
            
        student.save()
        
        current_admin = Admin.objects.get(admin_id=admin_id)
        AdminLog.objects.create(admin=current_admin, action=f"{log_action} student {student.username}")
        
        return JsonResponse({
            'success': True,
            'status': student.status,
            'message': f'Student {action} successfully'
        })
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=500)


@csrf_exempt
@require_http_methods(["POST"])
def edit_student(request, student_id):
    """Edit student details"""
    admin_id = request.session.get('user_id')
    if not admin_id:
        return JsonResponse({'error': 'Not authenticated'}, status=401)

    try:
        student = get_object_or_404(Student, pk=student_id)
        # Handle both JSON and Form data
        if request.content_type == 'application/json':
            data = json.loads(request.body)
        else:
            data = request.POST

        student.first_name = data.get('first_name', student.first_name)
        student.last_name = data.get('last_name', student.last_name)
        student.username = data.get('username', student.username)
        student.email = data.get('email', student.email)
        student.contact_num = data.get('contact_num', student.contact_num)
        student.address = data.get('address', student.address)
        student.barangay = data.get('barangay', student.barangay)
        student.student_type = data.get('student_type', student.student_type)
        
        # If explicitly updating to non-Undergraduate, clear program_and_yr
        if 'student_type' in data and data['student_type'] != 'Undergraduate':
            student.program_and_yr = None
        else:
            student.program_and_yr = data.get('program_and_yr', student.program_and_yr)
        
        # Only allow changing status if provided
        if 'status' in data:
            student.status = data['status']

        # Extended Profile Fields
        student.mname = data.get('mname', student.mname)
        student.elem_school = data.get('elem_school', student.elem_school)
        student.elem_year = data.get('elem_year', student.elem_year)
        student.jhs_school = data.get('jhs_school', student.jhs_school)
        student.jhs_year = data.get('jhs_year', student.jhs_year)
        student.shs_school = data.get('shs_school', student.shs_school)
        student.shs_year = data.get('shs_year', student.shs_year)
        student.college_school = data.get('college_school', student.college_school)
        student.college_year = data.get('college_year', student.college_year)
        student.achievements = data.get('achievements', student.achievements)
        student.parent_name = data.get('parent_name', student.parent_name)
        student.guardian_name = data.get('guardian_name', student.guardian_name)
        student.guardian_contact = data.get('guardian_contact', student.guardian_contact)

        student.save()

        current_admin = Admin.objects.get(admin_id=admin_id)
        AdminLog.objects.create(admin=current_admin, action=f"Edited student details for {student.username}")

        return JsonResponse({
            'success': True,
            'message': 'Student details updated successfully'
        })
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=500)


@csrf_exempt
@require_http_methods(["POST"])
def delete_student(request, student_id):
    """Delete a student account"""
    admin_id = request.session.get('user_id')
    if not admin_id:
        return JsonResponse({'error': 'Not authenticated'}, status=401)

    try:
        student = get_object_or_404(Student, pk=student_id)
        username = student.username
        student.delete()
        
        current_admin = Admin.objects.get(admin_id=admin_id)
        AdminLog.objects.create(admin=current_admin, action=f"Deleted student {username}")
        
        return JsonResponse({
            'success': True,
            'message': f'Student {username} deleted successfully'
        })
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=500)

@csrf_exempt
@require_http_methods(["POST"])
def update_student_profile(request):
    """Update student profile information"""
    student_id = request.session.get('user_id')
    if not student_id:
        return JsonResponse({'error': 'Not authenticated'}, status=401)
        
    try:
        student = get_object_or_404(Student, pk=student_id)
        if request.content_type == 'application/json':
            data = json.loads(request.body)
        else:
            data = request.POST

        student.first_name = data.get('first_name', student.first_name)
        student.last_name = data.get('last_name', student.last_name)
        
        bday_str = data.get('bday')
        if bday_str:
            student.bday = bday_str
            
        student.email = data.get('email', student.email)
        student.contact_num = data.get('contact_num', student.contact_num)
        student.address = data.get('address', student.address)
        student.barangay = data.get('barangay', student.barangay)
        student.student_type = data.get('student_type', student.student_type)
        
        # If explicitly updating to non-Undergraduate, clear program_and_yr
        if 'student_type' in data and data['student_type'] != 'Undergraduate':
            student.program_and_yr = None
        else:
            student.program_and_yr = data.get('program_and_yr', student.program_and_yr)
            
        student.sex = data.get('sex', student.sex)
        student.scholarship = data.get('scholarship', student.scholarship)
        
        # Extended Profile Fields
        student.mname = data.get('mname', student.mname)
        student.elem_school = data.get('elem_school', student.elem_school)
        student.elem_year = data.get('elem_year', student.elem_year)
        student.jhs_school = data.get('jhs_school', student.jhs_school)
        student.jhs_year = data.get('jhs_year', student.jhs_year)
        student.shs_school = data.get('shs_school', student.shs_school)
        student.shs_year = data.get('shs_year', student.shs_year)
        student.college_school = data.get('college_school', student.college_school)
        student.college_year = data.get('college_year', student.college_year)
        student.achievements = data.get('achievements', student.achievements)
        student.parent_name = data.get('parent_name', student.parent_name)
        student.guardian_name = data.get('guardian_name', student.guardian_name)
        student.guardian_contact = data.get('guardian_contact', student.guardian_contact)

        student.save()

        return JsonResponse({
            'success': True,
            'message': 'Profile updated successfully'
        })
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=500)


# Program Application Management Views
def get_program_applications(request):
    """Get all program applications for admin review"""
    admin_id = request.session.get('user_id')
    if not admin_id:
        return JsonResponse({'error': 'Not authenticated'}, status=401)
    
    try:
        # First check if we have any applications
        total_applications = Application.objects.count()
        print(f"Debug: Total applications found: {total_applications}")
        
        if total_applications == 0:
            return JsonResponse({'applications': []})
        
        # Get applications with related data
        applications = Application.objects.select_related('student', 'program').all().order_by('-app_id')
        application_data = []
        
        for app in applications:
            try:
                # Fetch global student documents
                student_docs_list = []
                for s_doc in app.student.documents.all():
                    student_docs_list.append({
                        'id': s_doc.id,
                        'name': s_doc.document_name,
                        'url': s_doc.file.url if s_doc.file else None
                    })

                # Safely access student data
                student_data = {
                    'username': getattr(app.student, 'username', 'Unknown'),
                    'first_name': getattr(app.student, 'first_name', 'Unknown'),
                    'last_name': getattr(app.student, 'last_name', 'Unknown'),
                    'email': getattr(app.student, 'email', 'Unknown'),
                    'doc_submitted': app.student.doc_submitted.url if app.student.doc_submitted else None,
                    'student_documents': student_docs_list,
                }
                
                # Fetch documents for this application
                documents_list = []
                app_docs = app.documents.all()
                for doc in app_docs:
                     documents_list.append({
                        'url': doc.file.url, 
                        'name': doc.file.name.split('/')[-1]
                     })
                
                # Safely access program data
                program_data = {
                    'program_name': getattr(app.program, 'program_name', 'Unknown Program'),
                }
                
                application_data.append({
                    'app_id': app.app_id,
                    'student': student_data,
                    'program': program_data,
                    'documents': documents_list,
                    'requirement_status': getattr(app, 'requirement_status', 'unknown'),
                    'remarks': getattr(app, 'remarks', ''),
                    'created_at': getattr(app, 'created_at', None)
                })
                
            except Exception as e:
                print(f"Debug: Error processing application {getattr(app, 'app_id', 'unknown')}: {str(e)}")
                continue
        
        return JsonResponse({'applications': application_data})
        
    except Exception as e:
        print(f"Debug: Error in get_program_applications: {str(e)}")
        import traceback
        traceback.print_exc()
        return JsonResponse({'error': f'Database error: {str(e)}'}, status=500)


@csrf_exempt
@require_http_methods(["POST"])
def approve_program_application(request, application_id):
    """Approve a program application"""
    admin_id = request.session.get('user_id')
    if not admin_id:
        return JsonResponse({'error': 'Not authenticated'}, status=401)
    
    try:
        data = json.loads(request.body)
        remarks = data.get('remarks', '')

        application = get_object_or_404(Application, app_id=application_id)
        application.requirement_status = 'approved'
        application.remarks = remarks
        application.save()
        
        current_admin = Admin.objects.get(admin_id=admin_id)
        AdminLog.objects.create(admin=current_admin, action=f"Approved program application {application.app_id}")
        
        # Send email notification
        try:
            student_email = application.student.email
            program_name = application.program.program_name
            subject = f"Application Approved: {program_name}"
            
            html_message = f"""
            <p>Dear {application.student.first_name},</p>
            <p>Congratulations! Your application for <strong>{program_name}</strong> has been approved.</p>
            <p><strong>Remarks from Admin:</strong></p>
            {remarks if remarks else '<p>No additional remarks.</p>'}
            <p>Thank you,</p>
            <p>ScholarSync Subic Team</p>
            """
            plain_message = strip_tags(html_message)
            
            send_mail(
                subject,
                plain_message,
                settings.DEFAULT_FROM_EMAIL,
                [student_email],
                html_message=html_message,
                fail_silently=True,
            )
        except Exception as e:
            print(f"Error sending email: {str(e)}")
        
        return JsonResponse({
            'success': True,
            'message': 'Program application approved successfully'
        })
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=500)


@csrf_exempt
@require_http_methods(["POST"])
def reject_program_application(request, application_id):
    """Reject a program application"""
    admin_id = request.session.get('user_id')
    if not admin_id:
        return JsonResponse({'error': 'Not authenticated'}, status=401)
    
    try:
        data = json.loads(request.body)
        remarks = data.get('remarks', '')

        application = get_object_or_404(Application, app_id=application_id)
        application.requirement_status = 'rejected'
        application.remarks = remarks
        application.save()
        
        current_admin = Admin.objects.get(admin_id=admin_id)
        AdminLog.objects.create(admin=current_admin, action=f"Rejected program application {application.app_id}")
        
        # Send email notification
        try:
            student_email = application.student.email
            program_name = application.program.program_name
            subject = f"Application Rejected: {program_name}"
            
            html_message = f"""
            <p>Dear {application.student.first_name},</p>
            <p>We regret to inform you that your application for <strong>{program_name}</strong> has been rejected.</p>
            <p><strong>Remarks from Admin:</strong></p>
            {remarks if remarks else '<p>No additional remarks.</p>'}
            <p>Thank you,</p>
            <p>ScholarSync Subic Team</p>
            """
            plain_message = strip_tags(html_message)
            
            send_mail(
                subject,
                plain_message,
                settings.DEFAULT_FROM_EMAIL,
                [student_email],
                html_message=html_message,
                fail_silently=True,
            )
        except Exception as e:
            print(f"Error sending email: {str(e)}")
        
        return JsonResponse({
            'success': True,
            'message': 'Program application rejected successfully'
        })
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=500)

@csrf_exempt
def get_program_applicants_by_program(request, program_id):
    """Get formatting applications for a specific program"""
    admin_id = request.session.get('user_id')
    if not admin_id:
        return JsonResponse({'error': 'Not authenticated'}, status=401)
        
    try:
        applications = Application.objects.filter(program_id=program_id).select_related('student', 'program').order_by('-created_at')
        
        app_list = []
        for app in applications:
            app_list.append({
                'app_id': app.app_id,
                'student_name': f"{app.student.first_name} {app.student.last_name}",
                'username': app.student.username,
                'requirement_status': app.requirement_status,
                'created_at': app.created_at.strftime('%Y-%m-%d %H:%M') if app.created_at else None,
                'program_type': app.program.program_type
            })
            
        return JsonResponse({'success': True, 'applications': app_list})
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=500)

# Chart Data API Views
def get_application_trends(request):
    """Get application trends data for charts"""
    admin_id = request.session.get('user_id')
    if not admin_id:
        return JsonResponse({'error': 'Not authenticated'}, status=401)
    
    try:
        period = request.GET.get('period', '30')  # Default to 30 days
        days = int(period)
        
        # Calculate start date
        end_date = timezone.now().date()
        start_date = end_date - timedelta(days=days-1)
        
        # Get application data grouped by day
        applications_by_day = Application.objects.filter(
            created_at__date__gte=start_date,
            created_at__date__lte=end_date
        ).extra(
            select={'day': 'date(created_at)'}
        ).values('day').annotate(count=Count('app_id')).order_by('day')
        
        # Create a complete date range with counts
        date_counts = {}
        for item in applications_by_day:
            date_counts[item['day']] = item['count']
        
        # Generate labels and data arrays
        labels = []
        data = []
        current_date = start_date
        
        while current_date <= end_date:
            date_str = current_date.strftime('%b %d')
            labels.append(date_str)
            data.append(date_counts.get(current_date, 0))
            current_date += timedelta(days=1)
        
        return JsonResponse({
            'labels': labels,
            'data': data
        })
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)


def get_student_statistics(request):
    """Get student statistics data for charts"""
    admin_id = request.session.get('user_id')
    if not admin_id:
        return JsonResponse({'error': 'Not authenticated'}, status=401)
    
    try:
        chart_type = request.GET.get('type', 'status')
        
        if chart_type == 'status':
            # Get students by status
            status_data = Student.objects.values('status').annotate(count=Count('id'))
            
            labels = []
            data = []
            colors = []
            
            status_colors = {
                'active': 'rgba(102, 126, 234, 0.8)',
                'pending': 'rgba(243, 156, 18, 0.8)',
                'rejected': 'rgba(231, 76, 60, 0.8)'
            }
            
            for item in status_data:
                labels.append(item['status'].title())
                data.append(item['count'])
                colors.append(status_colors.get(item['status'], 'rgba(102, 126, 234, 0.8)'))
            
        elif chart_type == 'program':
            # Get students by program
            program_data = Student.objects.values('program_and_yr').annotate(count=Count('id')).order_by('-count')[:10]
            
            labels = []
            data = []
            colors = [
                'rgba(102, 126, 234, 0.8)',
                'rgba(118, 75, 162, 0.8)',
                'rgba(39, 174, 96, 0.8)',
                'rgba(243, 156, 18, 0.8)',
                'rgba(231, 76, 60, 0.8)',
                'rgba(52, 152, 219, 0.8)',
                'rgba(155, 89, 182, 0.8)',
                'rgba(46, 204, 113, 0.8)',
                'rgba(230, 126, 34, 0.8)',
                'rgba(231, 76, 60, 0.8)'
            ]
            
            for i, item in enumerate(program_data):
                labels.append(item['program_and_yr'])
                data.append(item['count'])
            
            colors = colors[:len(labels)]
            
        elif chart_type == 'monthly':
            # Get monthly registrations for the last 6 months
            end_date = timezone.now().date()
            start_date = end_date - timedelta(days=180)  # 6 months
            
            monthly_data = Student.objects.filter(
                created_at__date__gte=start_date
            ).extra(
                select={'month': 'strftime("%Y-%m", created_at)'}
            ).values('month').annotate(count=Count('id')).order_by('month')
            
            labels = []
            data = []
            
            for item in monthly_data:
                # Format month label
                month_date = datetime.strptime(item['month'], '%Y-%m')
                labels.append(month_date.strftime('%b'))
                data.append(item['count'])
        
        return JsonResponse({
            'labels': labels,
            'data': data,
            'colors': colors if chart_type in ['status', 'program'] else None
        })
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)


@csrf_exempt
def get_my_applications(request):
    """Get logged-in student's applications"""
    student_id = request.session.get('user_id')
    if not student_id:
        return JsonResponse({'error': 'Not authenticated'}, status=401)

    try:
        student = get_object_or_404(Student, pk=student_id)
        applications = Application.objects.filter(student=student).select_related('program').order_by('-created_at')
        
        app_list = []
        for app in applications:
            app_list.append({
                'app_id': app.app_id,
                'program_name': app.program.program_name,
                'program_image': app.program.program_image.url if app.program.program_image else None,
                'requirement_status': app.requirement_status,
                'remarks': app.remarks,
                'is_remarks_viewed': app.is_remarks_viewed,
                'created_at': app.created_at.isoformat() if app.created_at else None
            })
            
        return JsonResponse({'applications': app_list})
    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)


@csrf_exempt
@require_http_methods(["POST"])
def mark_remarks_viewed(request, application_id):
    """Mark application remarks as viewed"""
    student_id = request.session.get('user_id')
    if not student_id:
        return JsonResponse({'error': 'Not authenticated'}, status=401)
        
    try:
        # Verify the application belongs to this student
        app = get_object_or_404(Application, app_id=application_id, student__pk=student_id)
        
        app.is_remarks_viewed = True
        app.save()
        
        return JsonResponse({'success': True})
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=500)


def student_voucher_view(request, application_id):
    """View to generic voucher for accepted students"""
    student_id = request.session.get('user_id')
    if not student_id:
        return redirect('accounts:login')
    
    student = get_object_or_404(Student, pk=student_id)
    application = get_object_or_404(Application, app_id=application_id, student=student)
    
    # Check if application is approved
    if application.requirement_status != 'approved':
        # You might want to show a message or redirect
        messages.error(request, "Voucher is only available for approved applications.")
        return redirect('accounts:student_dashboard')
        
    context = {
        'student': student,
        'application': application,
        'program': application.program,
        'now': timezone.now()
    }
    return render(request, 'accounts/student_voucher.html', context)


def admin_receipt_view(request, application_id):
    """View to print receiving copy / receipt for accepted students in Financial Assistance"""
    admin_id = request.session.get('user_id')
    if not admin_id:
        return redirect('accounts:login')
    
    application = get_object_or_404(Application, app_id=application_id)
    student = application.student
    
    amount = request.GET.get('amount', '0')
    
    # Check if application is approved
    if application.requirement_status != 'approved':
        return HttpResponse("Receipt is only available for approved applications.", status=400)
    
    if application.program.program_type != 'Financial Assistance':
        return HttpResponse("Receipt corresponds to Financial Assistance only.", status=400)
        
    # Format amount with commas
    try:
        amount_float = float(amount.replace(',', ''))
        formatted_amount = f"{amount_float:,.2f}"
    except ValueError:
        formatted_amount = amount

    context = {
        'student': student,
        'application': application,
        'program': application.program,
        'now': timezone.now(),
        'amount': formatted_amount
    }
    return render(request, 'accounts/admin_receipt.html', context)
    

def generate_report(request):
    """Generate and export reports"""
    admin_id = request.session.get('user_id')
    if not admin_id:
        return JsonResponse({'error': 'Not authenticated'}, status=401)

    try:
        report_type = request.GET.get('type')
        start_date = request.GET.get('start_date')
        end_date = request.GET.get('end_date')
        status = request.GET.get('status')
        program_id = request.GET.get('program')
        export_csv = request.GET.get('export') == 'csv'

        data = []
        filename = f"report_{report_type}_{datetime.now().strftime('%Y%m%d')}"
        
        # Log generation event if exporting
        if request.GET.get('export'):
            export_format = "Word" if request.GET.get('export') == 'word' else "CSV"
            current_admin = Admin.objects.get(admin_id=admin_id)
            AdminLog.objects.create(admin=current_admin, action=f"Generated {report_type} report in {export_format} format")

        if report_type == 'students':
            # Base query
            query = Student.objects.all().order_by('-created_at')

            # Filters
            if start_date:
                query = query.filter(created_at__date__gte=start_date)
            if end_date:
                query = query.filter(created_at__date__lte=end_date)
            if status and status != 'all':
                query = query.filter(status=status)
            
            barangay = request.GET.get('barangay')
            school = request.GET.get('school')
            
            if barangay and barangay != 'all':
                query = query.filter(barangay__iexact=barangay)
            if school:
                query = query.filter(current_school__icontains=school)
            
            # For data creation
            header = ['Student ID', 'Username', 'First Name', 'Last Name', 'Email', 'Program', 'Barangay', 'School', 'Status', 'Date Joined']
            
            for s in query:
                data.append({
                    'id': s.student_id,
                    'username': s.username,
                    'first_name': s.first_name,
                    'last_name': s.last_name,
                    'email': s.email,
                    'program': s.program_and_yr,
                    'barangay': s.barangay,
                    'school': s.current_school,
                    'status': s.status,
                    'created_at': s.created_at.strftime('%Y-%m-%d') if s.created_at else 'N/A'
                })

        elif report_type == 'applications':
            # Base query
            query = Application.objects.select_related('student', 'program').all().order_by('-created_at')

            # Filters
            if start_date:
                query = query.filter(created_at__date__gte=start_date)
            if end_date:
                query = query.filter(created_at__date__lte=end_date)
            if status and status != 'all':
                query = query.filter(requirement_status=status)
            if program_id and program_id != 'all':
                query = query.filter(program_id=program_id)
                
            barangay = request.GET.get('barangay')
            school = request.GET.get('school')
            
            if barangay and barangay != 'all':
                query = query.filter(student__barangay__iexact=barangay)
            if school:
                query = query.filter(student__current_school__icontains=school)

            header = ['App ID', 'Student', 'Program', 'Barangay', 'School', 'Status', 'Date Submitted']
            
            for app in query:
                data.append({
                    'id': app.app_id,
                    'student': f"{app.student.first_name} {app.student.last_name}",
                    'program': app.program.program_name,
                    'barangay': app.student.barangay,
                    'school': app.student.current_school,
                    'status': app.requirement_status,
                    'created_at': app.created_at.strftime('%Y-%m-%d') if app.created_at else 'N/A'
                })
        
        else:
             return JsonResponse({'error': 'Invalid report type'}, status=400)

        # Handle CSV Export
        if export_csv:
            response = HttpResponse(content_type='text/csv')
            response['Content-Disposition'] = f'attachment; filename="{filename}.csv"'

            fields_param = request.GET.get('fields')
            selected_fields = fields_param.split(',') if fields_param else ['id', 'username', 'name', 'email', 'program', 'status', 'date', 'barangay', 'school']

            writer = csv.writer(response)

            if report_type == 'students':
                export_header = []
                if 'id' in selected_fields: export_header.append('Student ID')
                if 'username' in selected_fields: export_header.append('Username')
                if 'name' in selected_fields: export_header.extend(['First Name', 'Last Name'])
                if 'email' in selected_fields: export_header.append('Email')
                if 'program' in selected_fields: export_header.append('Program')
                if 'barangay' in selected_fields: export_header.append('Barangay')
                if 'school' in selected_fields: export_header.append('School')
                if 'status' in selected_fields: export_header.append('Status')
                if 'date' in selected_fields: export_header.append('Date Joined')
                
                writer.writerow(export_header)
                
                for row in data:
                    out_row = []
                    if 'id' in selected_fields: out_row.append(row['id'])
                    if 'username' in selected_fields: out_row.append(row['username'])
                    if 'name' in selected_fields: out_row.extend([row['first_name'], row['last_name']])
                    if 'email' in selected_fields: out_row.append(row['email'])
                    if 'program' in selected_fields: out_row.append(row['program'])
                    if 'barangay' in selected_fields: out_row.append(row['barangay'])
                    if 'school' in selected_fields: out_row.append(row['school'])
                    if 'status' in selected_fields: out_row.append(row['status'])
                    if 'date' in selected_fields: out_row.append(row['created_at'])
                    writer.writerow(out_row)
            else:
                export_header = []
                if 'id' in selected_fields: export_header.append('App ID')
                if 'name' in selected_fields: export_header.append('Student')
                if 'program' in selected_fields: export_header.append('Program')
                if 'barangay' in selected_fields: export_header.append('Barangay')
                if 'school' in selected_fields: export_header.append('School')
                if 'status' in selected_fields: export_header.append('Status')
                if 'date' in selected_fields: export_header.append('Date Submitted')
                
                writer.writerow(export_header)
                
                for row in data:
                    out_row = []
                    if 'id' in selected_fields: out_row.append(row['id'])
                    if 'name' in selected_fields: out_row.append(row['student'])
                    if 'program' in selected_fields: out_row.append(row['program'])
                    if 'barangay' in selected_fields: out_row.append(row['barangay'])
                    if 'school' in selected_fields: out_row.append(row['school'])
                    if 'status' in selected_fields: out_row.append(row['status'])
                    if 'date' in selected_fields: out_row.append(row['created_at'])
                    writer.writerow(out_row)

            return response
        
        # Handle Word Export
        # Handle Word Export
        if request.GET.get('export') == 'word':
            document = Document()
            
            # Set Font to Calibri
            style = document.styles['Normal']
            font = style.font
            font.name = 'Calibri'
            font.size = Pt(11)
            
            # Adjust Margins
            section = document.sections[0]
            section.top_margin = Inches(0.5)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            
            # Header Setup
            header_section = section.header
            header_section.is_linked_to_previous = False
            
            # Use a table for header layout to align Logo Left and Title Center
            # 3 columns: Left (Logo), Center (Title), Right (Logo)
            htable = header_section.add_table(1, 3, width=Inches(7.5)) # Width = 8.5 - 0.5 - 0.5
            htable.autofit = False
            htable.columns[0].width = Inches(1.5)
            htable.columns[1].width = Inches(4.5)
            htable.columns[2].width = Inches(1.5)
            
            # Cell 0: Logo
            cell0 = htable.cell(0, 0)
            p0 = cell0.paragraphs[0]
            p0.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            logo_path = os.path.join(settings.BASE_DIR, 'accounts', 'static', 'accounts', 'subic_seal.png')
            if os.path.exists(logo_path):
                run0 = p0.add_run()
                run0.add_picture(logo_path, height=Inches(0.8))
            
            # Cell 1: Title and Address
            cell1 = htable.cell(0, 1)
            p1 = cell1.paragraphs[0]
            p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run1 = p1.add_run('ScholarSync Subic\n')
            run1.bold = True
            run1.font.size = Pt(16) 
            
            run_address = p1.add_run('Baraca-Camachile, National Highway, Subic, 2209 Zambales\n')
            run_address.font.size = Pt(10)
            
            run_title = p1.add_run('Applicants List')
            run_title.bold = True
            run_title.font.size = Pt(14)
            
            # Cell 2: Second Logo
            cell2 = htable.cell(0, 2)
            p2 = cell2.paragraphs[0]
            p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            logo2_path = os.path.join(settings.BASE_DIR, 'accounts', 'static', 'accounts', 'scholarsync_logo.png')
            if os.path.exists(logo2_path):
                run2 = p2.add_run()
                run2.add_picture(logo2_path, height=Inches(0.8))
            
            # Remove the default empty paragraph in the header if it causes spacing issues
            if len(header_section.paragraphs) > 0:
                header_section.paragraphs[0]._element.getparent().remove(header_section.paragraphs[0]._element)

            # Add a paragraph for the separator line
            border_paragraph = header_section.add_paragraph()
            border_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            border_paragraph.paragraph_format.space_after = Pt(6)
            
            # Apply bottom border to this paragraph using OXML
            p = border_paragraph._p
            pPr = p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '6')  # 1/8 pt, so 6 is 3/4 pt. Thin line.
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), 'CCCCCC') # Gray color
            pBdr.append(bottom)
            pPr.append(pBdr)

            # Map fields for Word export
            fields_param = request.GET.get('fields')
            selected_fields = fields_param.split(',') if fields_param else ['id', 'username', 'name', 'email', 'program', 'status', 'date', 'barangay', 'school']

            if report_type == 'students':
                word_header = []
                if 'id' in selected_fields: word_header.append('Student ID')
                if 'username' in selected_fields: word_header.append('Username')
                if 'name' in selected_fields: word_header.extend(['First Name', 'Last Name'])
                if 'email' in selected_fields: word_header.append('Email')
                if 'program' in selected_fields: word_header.append('Program')
                if 'barangay' in selected_fields: word_header.append('Barangay')
                if 'school' in selected_fields: word_header.append('School')
                if 'status' in selected_fields: word_header.append('Status')
                if 'date' in selected_fields: word_header.append('Date Joined')
            else:
                word_header = []
                if 'id' in selected_fields: word_header.append('App ID')
                if 'name' in selected_fields: word_header.append('Student')
                if 'program' in selected_fields: word_header.append('Program')
                if 'barangay' in selected_fields: word_header.append('Barangay')
                if 'school' in selected_fields: word_header.append('School')
                if 'status' in selected_fields: word_header.append('Status')
                if 'date' in selected_fields: word_header.append('Date Submitted')

            # Build data rows correspondingly
            rows_data = []
            if report_type == 'students':
                for row in data:
                    out_row = []
                    if 'id' in selected_fields: out_row.append(str(row['id']))
                    if 'username' in selected_fields: out_row.append(str(row['username']))
                    if 'name' in selected_fields: out_row.extend([str(row['first_name']), str(row['last_name'])])
                    if 'email' in selected_fields: out_row.append(str(row['email']))
                    if 'program' in selected_fields: out_row.append(str(row['program']))
                    if 'barangay' in selected_fields: out_row.append(str(row['barangay']))
                    if 'school' in selected_fields: out_row.append(str(row['school']))
                    if 'status' in selected_fields: out_row.append(str(row['status']))
                    if 'date' in selected_fields: out_row.append(str(row.get('created_at', '')))
                    rows_data.append(out_row)
            else:
                for row in data:
                    out_row = []
                    if 'id' in selected_fields: out_row.append(str(row['id']))
                    if 'name' in selected_fields: out_row.append(str(row['student']))
                    if 'program' in selected_fields: out_row.append(str(row['program']))
                    if 'barangay' in selected_fields: out_row.append(str(row['barangay']))
                    if 'school' in selected_fields: out_row.append(str(row['school']))
                    if 'status' in selected_fields: out_row.append(str(row['status']))
                    if 'date' in selected_fields: out_row.append(str(row.get('created_at', '')))
                    rows_data.append(out_row)

            # Table
            table = document.add_table(rows=1, cols=len(word_header))
            table.style = 'Table Grid'
            
            # Helper function for shading
            def set_cell_background(cell, color_hex):
                tcPr = cell._tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), color_hex)
                tcPr.append(shd)

            # Table Header
            hdr_cells = table.rows[0].cells
            for i, col_name in enumerate(word_header):
                cell = hdr_cells[i]
                set_cell_background(cell, '2E8B57') # Sea Green
                run = cell.paragraphs[0].add_run(col_name)
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255) # White text
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Table Data Content
            for index, row_content in enumerate(rows_data):
                row_cells = table.add_row().cells
                is_even = (index % 2 == 0)
                bg_color = 'E8F5E9' if is_even else 'FFFFFF' # Light Green / White
                
                for i, text in enumerate(row_content):
                    if i < len(row_cells):
                        cell = row_cells[i]
                        cell.text = text
                        if is_even: 
                             set_cell_background(cell, bg_color)
            
            # Footer Setup (Metadata, Page Number, Approved By)
            footer = section.footer
            footer.is_linked_to_previous = False
            
            ftable = footer.add_table(1, 3, width=Inches(7.5))
            ftable.autofit = False
            ftable.columns[0].width = Inches(2.5)
            ftable.columns[1].width = Inches(2.5)
            ftable.columns[2].width = Inches(2.5)
            
            # Left Footer: Generated On
            fcell0 = ftable.cell(0, 0)
            fp0 = fcell0.paragraphs[0]
            fp0.alignment = WD_ALIGN_PARAGRAPH.LEFT
            fr0 = fp0.add_run(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
            fr0.font.color.rgb = RGBColor(128, 128, 128)
            fr0.font.size = Pt(9)
            if start_date or end_date:
                sd = start_date if start_date else 'Any'
                ed = end_date if end_date else 'Any'
                fr_dates = fp0.add_run(f'\nPeriod: {sd} to {ed}')
                fr_dates.font.color.rgb = RGBColor(128, 128, 128)
                fr_dates.font.size = Pt(9)
            
            # Middle Footer: Page Numbers
            fcell1 = ftable.cell(0, 1)
            fp1 = fcell1.paragraphs[0]
            fp1.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            def add_page_number(run):
                fldChar1 = OxmlElement('w:fldChar')
                fldChar1.set(qn('w:fldCharType'), 'begin')
                instrText = OxmlElement('w:instrText')
                instrText.set(qn('xml:space'), 'preserve')
                instrText.text = "PAGE"
                fldChar2 = OxmlElement('w:fldChar')
                fldChar2.set(qn('w:fldCharType'), 'separate')
                fldChar3 = OxmlElement('w:fldChar')
                fldChar3.set(qn('w:fldCharType'), 'end')
                run._r.append(fldChar1)
                run._r.append(instrText)
                run._r.append(fldChar2)
                run._r.append(fldChar3)
                
            fr1 = fp1.add_run('Page ')
            fr1.font.color.rgb = RGBColor(128, 128, 128)
            fr1.font.size = Pt(9)
            add_page_number(fp1.add_run())

            # Right Footer: Approved By
            fcell2 = ftable.cell(0, 2)
            fp2 = fcell2.paragraphs[0]
            fp2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            fr2 = fp2.add_run('Approved by: ____________________')
            fr2.font.color.rgb = RGBColor(0, 0, 0)
            fr2.font.size = Pt(11)

            response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            response['Content-Disposition'] = f'attachment; filename="{filename}.docx"'
            document.save(response)
            return response
        
        # Return JSON for table view
        return JsonResponse({'data': data})

    except Exception as e:
        return JsonResponse({'error': str(e)}, status=500)

# -------------------------------------------------------------
# Messaging and Ticketing System Views
# -------------------------------------------------------------

@csrf_exempt
@require_http_methods(["GET"])
def get_messages(request):
    """Fetch messages for the current user (admin or student)."""
    user_id = request.session.get('user_id')
    user_role = request.session.get('user_role')

    if not user_id:
        return JsonResponse({'error': 'Not authenticated'}, status=401)

    try:
        if user_role == 'admin':
            messages_qs = Message.objects.all().order_by('-created_at')
        else:
            student = get_object_or_404(Student, pk=user_id)
            messages_qs = Message.objects.filter(student=student).order_by('-created_at')

        messages_data = []
        for msg in messages_qs:
            admin_name = msg.admin.admin_name if msg.admin else 'System/Unknown'
            student_name = f"{msg.student.first_name} {msg.student.last_name} ({msg.student.username})"
            
            messages_data.append({
                'id': msg.id,
                'sender_type': msg.sender_type,
                'sender_name': admin_name if msg.sender_type == 'admin' else student_name,
                'target_student_id': msg.student.student_id,
                'target_student_name': student_name,
                'subject': msg.subject,
                'body': msg.body,
                'is_read': msg.is_read,
                'created_at': msg.created_at.isoformat()
            })

        return JsonResponse({'success': True, 'messages': messages_data})
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=500)


@csrf_exempt
@require_http_methods(["POST"])
def send_message(request):
    """Send a new message or ticket."""
    user_id = request.session.get('user_id')
    user_role = request.session.get('user_role')

    if not user_id:
        return JsonResponse({'error': 'Not authenticated'}, status=401)

    try:
        data = json.loads(request.body)
        subject = data.get('subject')
        body = data.get('body')
        
        if not subject or not body:
            return JsonResponse({'success': False, 'error': 'Subject and body are required.'}, status=400)

        if user_role == 'admin':
            admin = get_object_or_404(Admin, pk=user_id)
            student_id = data.get('student_id')
            if not student_id:
                return JsonResponse({'success': False, 'error': 'Student ID is required for admins to send a message.'}, status=400)
            
            student = get_object_or_404(Student, pk=student_id)
            Message.objects.create(
                student=student,
                admin=admin,
                sender_type='admin',
                subject=subject,
                body=body
            )
            
            AdminLog.objects.create(admin=admin, action=f"Sent message to student {student.username}")
        else:
            # Student is submitting a ticket
            student = get_object_or_404(Student, pk=user_id)
            Message.objects.create(
                student=student,
                admin=None, # TBD: Might assign to generic admin or specific admin if requested
                sender_type='student',
                subject=subject,
                body=body
            )

        return JsonResponse({'success': True, 'message': 'Message sent successfully.'})
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=500)


@csrf_exempt
@require_http_methods(["POST"])
def admin_send_system_message(request):
    """Send a direct system message to a student."""
    user_id = request.session.get('user_id')
    user_role = request.session.get('user_role')

    if not user_id or user_role != 'admin':
        return JsonResponse({'success': False, 'error': 'Not authorized'}, status=403)

    try:
        admin = get_object_or_404(Admin, pk=user_id)
        student_id = request.POST.get('student_id')
        subject = request.POST.get('subject')
        body = request.POST.get('body')

        if not student_id or not subject or not body:
            return JsonResponse({'success': False, 'error': 'Missing required fields.'}, status=400)

        student = get_object_or_404(Student, pk=student_id)
        
        Message.objects.create(
            student=student,
            admin=admin,
            sender_type='admin',
            subject=subject,
            body=body
        )

        AdminLog.objects.create(admin=admin, action=f"Sent direct system message to {student.username}")
        return JsonResponse({'success': True})
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=500)


@csrf_exempt
@require_http_methods(["POST"])
def admin_send_email(request):
    """Send an email wrapper to a student, with optional attachment."""
    user_id = request.session.get('user_id')
    user_role = request.session.get('user_role')

    if not user_id or user_role != 'admin':
        return JsonResponse({'success': False, 'error': 'Not authorized'}, status=403)

    try:
        admin = get_object_or_404(Admin, pk=user_id)
        student_id = request.POST.get('student_id')
        subject = request.POST.get('subject')
        body_html = request.POST.get('body')
        attachment = request.FILES.get('attachment')

        if not student_id or not subject or not body_html:
            return JsonResponse({'success': False, 'error': 'Missing required fields.'}, status=400)

        student = get_object_or_404(Student, pk=student_id)
        
        if not student.email:
            return JsonResponse({'success': False, 'error': 'Student does not have an email address.'}, status=400)

        # Build Email
        from_email = settings.EMAIL_HOST_USER if hasattr(settings, 'EMAIL_HOST_USER') and settings.EMAIL_HOST_USER else 'admin@scholarsync.com'
        text_content = strip_tags(body_html)
        
        msg = EmailMultiAlternatives(subject, text_content, from_email, [student.email])
        msg.attach_alternative(body_html, "text/html")
        
        if attachment:
            msg.attach(attachment.name, attachment.read(), attachment.content_type)
            
        # Send Email
        def send_email_thread():
            msg.send(fail_silently=True)
            
        threading.Thread(target=send_email_thread).start()

        AdminLog.objects.create(admin=admin, action=f"Sent direct email to {student.username}")
        return JsonResponse({'success': True})
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=500)


@csrf_exempt
@require_http_methods(["POST"])
def admin_send_batch_message(request):
    """Send a batch message (email or system) to filtered students."""
    user_id = request.session.get('user_id')
    user_role = request.session.get('user_role')

    if not user_id or user_role != 'admin':
        return JsonResponse({'success': False, 'error': 'Not authorized'}, status=403)

    try:
        admin = get_object_or_404(Admin, pk=user_id)
        
        mode = request.POST.get('mode') # 'email' or 'system'
        status_filter = request.POST.get('status', 'all')
        program_filter = request.POST.get('program', 'all')
        type_filter = request.POST.get('student_type', 'all')
        barangay_filter = request.POST.get('barangay', 'all')
        school_filter = request.POST.get('school', '')
        
        subject = request.POST.get('subject')
        body_html = request.POST.get('body')
        attachment = request.FILES.get('attachment')

        if not mode or not subject or not body_html:
            return JsonResponse({'success': False, 'error': 'Missing required fields.'}, status=400)

        # Build queryset based on filters (same logic as report generation)
        students = Student.objects.all()
        
        if status_filter != 'all':
            students = students.filter(status=status_filter)
            
        if type_filter != 'all':
            students = students.filter(student_type=type_filter)
            
        if program_filter != 'all':
            # Need to join with applications to filter by program
            students = students.filter(applications__program__program_name__icontains=program_filter).distinct()
            
        if barangay_filter != 'all':
            students = students.filter(barangay__iexact=barangay_filter)
            
        if school_filter:
            students = students.filter(current_school__icontains=school_filter)

        student_count = students.count()
        if student_count == 0:
            return JsonResponse({'success': False, 'error': 'No students matched the selected filters.'}, status=400)

        if mode == 'system':
            # Create a system message for each student
            messages_to_create = [
                Message(
                    student=student,
                    admin=admin,
                    sender_type='admin',
                    subject=subject,
                    body=body_html
                ) for student in students
            ]
            Message.objects.bulk_create(messages_to_create)
            AdminLog.objects.create(admin=admin, action=f"Sent batch system message to {student_count} students")

        elif mode == 'email':
            # Send an email to each student
            from_email = settings.EMAIL_HOST_USER if hasattr(settings, 'EMAIL_HOST_USER') and settings.EMAIL_HOST_USER else 'admin@scholarsync.com'
            text_content = strip_tags(body_html)
            
            attachment_data = None
            if attachment:
                attachment_data = {
                    'name': attachment.name,
                    'content': attachment.read(),
                    'type': attachment.content_type
                }
            
            # Send emails in a background thread to prevent blocking
            def send_batch_emails_thread(student_list, att_data):
                for student in student_list:
                    if not student.email:
                        continue
                    
                    try:
                        msg = EmailMultiAlternatives(subject, text_content, from_email, [student.email])
                        msg.attach_alternative(body_html, "text/html")
                        
                        if att_data:
                            msg.attach(att_data['name'], att_data['content'], att_data['type'])
                            
                        msg.send(fail_silently=True)
                    except Exception as e:
                        print(f"Failed to send email to {student.email}: {e}")

            threading.Thread(target=send_batch_emails_thread, args=(list(students), attachment_data)).start()
            AdminLog.objects.create(admin=admin, action=f"Sent batch email to {student_count} students")

        return JsonResponse({'success': True, 'count': student_count})
        
    except Exception as e:
        import traceback
        traceback.print_exc()
        return JsonResponse({'success': False, 'error': str(e)}, status=500)



# Admin: get a specific student's documents
@csrf_exempt
@require_http_methods(["GET"])
def admin_get_student_documents(request, student_id):
    if not request.session.get('user_id'):
        return JsonResponse({'error': 'Not authenticated'}, status=401)
    try:
        from home.models import Program as HomeProgram
        student = get_object_or_404(Student, pk=student_id)

        # Uploaded documents
        docs = student.documents.all().order_by('-uploaded_at')
        uploaded = [
            {'id': d.id, 'name': d.document_name, 'url': d.file.url if d.file else None,
             'uploaded_at': d.uploaded_at.isoformat()}
            for d in docs
        ]

        # Standard student dashboard checklist 
        required_set = [
            "School ID (Current Semester)",
            "Valid ID",
            "TOR",
            "Voter's Certificate",
            "Certificate of Indigency",
            "Letter of Application"
        ]
        seen = set(required_set)
        
        # Also collect any unique required document names from programs just in case
        for program in HomeProgram.objects.all():
            for req in (program.document_requirements or []):
                if req and req not in seen:
                    seen.add(req)
                    required_set.append(req)

        # Build uploaded names set for quick lookup
        uploaded_names = {d['name'] for d in uploaded if d['name']}

        return JsonResponse({
            'success': True,
            'documents': uploaded,
            'required_documents': required_set,
            'uploaded_names': list(uploaded_names),
        })
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=500)



# -------------------------------------------------------------
# Document Management Views (Student)
# -------------------------------------------------------------

@csrf_exempt
@require_http_methods(["GET"])
def get_student_documents(request):
    """Fetch documents for the logged in student"""
    student_id = request.session.get('user_id')
    if not student_id:
        return JsonResponse({'error': 'Not authenticated'}, status=401)
        
    try:
        student = get_object_or_404(Student, pk=student_id)
        documents = student.documents.all().order_by('-uploaded_at')
        doc_list = []
        for doc in documents:
            doc_list.append({
                'id': doc.id,
                'name': doc.document_name,
                'url': doc.file.url if doc.file else None,
                'uploaded_at': doc.uploaded_at.isoformat()
            })
        return JsonResponse({'success': True, 'documents': doc_list})
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=500)


@csrf_exempt
@require_http_methods(["POST"])
def upload_student_document(request):
    """Upload a new document for the logged in student"""
    student_id = request.session.get('user_id')
    if not student_id:
        return JsonResponse({'error': 'Not authenticated'}, status=401)
        
    try:
        student = get_object_or_404(Student, pk=student_id)
        
        # Check if the student has reached the maximum document limit (15)
        if student.documents.count() >= 15:
            return JsonResponse({'success': False, 'error': 'Maximum limit of 15 documents reached. Please delete some documents to upload new ones.'}, status=400)
            
        document_name = request.POST.get('document_name')
        file = request.FILES.get('file')
        
        if not document_name or not file:
             return JsonResponse({'success': False, 'error': 'Document name and file are required'}, status=400)
             
        doc = StudentDocument.objects.create(
            student=student,
            document_name=document_name,
            file=file
        )
        return JsonResponse({'success': True, 'message': 'Document uploaded successfully', 'doc_id': doc.id})
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=500)


@csrf_exempt
@require_http_methods(["POST", "DELETE"])
def delete_student_document(request, doc_id):
    """Delete a student document"""
    student_id = request.session.get('user_id')
    if not student_id:
        return JsonResponse({'error': 'Not authenticated'}, status=401)
        
    try:
        doc = get_object_or_404(StudentDocument, id=doc_id, student__pk=student_id)
        doc.delete()
        return JsonResponse({'success': True, 'message': 'Document deleted successfully'})
    except Exception as e:
         return JsonResponse({'success': False, 'error': str(e)}, status=500)


@csrf_exempt
@require_http_methods(["POST"])
def mark_message_read(request, message_id):
    """Mark a message as read."""
    user_id = request.session.get('user_id')
    user_role = request.session.get('user_role')

    if not user_id:
        return JsonResponse({'error': 'Not authenticated'}, status=401)

    try:
        msg = get_object_or_404(Message, id=message_id)
        
        # Security: ensure only the intended recipient or an admin can mark it read
        if user_role == 'student' and msg.student.student_id != user_id:
            return JsonResponse({'success': False, 'error': 'Unauthorized'}, status=403)
            
        msg.is_read = True
        msg.save()
        return JsonResponse({'success': True})
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=500)

@csrf_exempt
@require_http_methods(["POST"])
def student_change_password(request):
    """Change student password"""
    student_id = request.session.get('user_id')
    if not student_id:
        return JsonResponse({'error': 'Not authenticated'}, status=401)
    
    try:
        data = json.loads(request.body)
        current_password = data.get('current_password')
        new_password = data.get('new_password')
        confirm_password = data.get('confirm_password')
        
        if not current_password or not new_password or not confirm_password:
            return JsonResponse({'success': False, 'error': 'All password fields are required.'})
            
        student = get_object_or_404(Student, pk=student_id)
        
        if student.password != current_password:
            return JsonResponse({'success': False, 'error': 'Incorrect current password.'})
            
        if new_password != confirm_password:
            return JsonResponse({'success': False, 'error': 'New passwords do not match.'})
            
        # Password Validation
        if not re.match(r'^(?=.*\d)(?=.*[A-Z]).{8,}$', new_password):
            return JsonResponse({'success': False, 'error': 'Password must be at least 8 characters long, contain at least one capital letter and one number.'})
            
        student.password = new_password
        student.save()
        
        return JsonResponse({'success': True, 'message': 'Password changed successfully.'})
        
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=500)
